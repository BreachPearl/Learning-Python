from docx import Document
with open("C:/Learning Python/ReadnWrite/Invited_names.txt") as file2:
    names_list=file2.readlines()

birtdayletter_path = "C:/Learning Python/ReadnWrite/Birthday.docx"
birthdayletter = Document(birtdayletter_path)
    

for name in names_list:
    stripped_name=name.strip("\n")
    new_document=Document()
    for paragraph in birthdayletter.paragraphs:
        new_paragraph=new_document.add_paragraph()
        new_paragraph.add_run(paragraph.text.replace("Name",stripped_name))

    new_document.save(f"C:/Learning Python/ReadnWrite/birthday_letters/{stripped_name}_letter.docx")
    
